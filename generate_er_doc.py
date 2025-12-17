import os
import subprocess
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from eralchemy2 import render_er


def create_connection_string(host, user, password, db_name, port=None):
    port = port or 5432
    return f"postgresql+psycopg2://{user}:{password}@{host}:{port}/{db_name}"


def read_body_content(filepath):
    project = "Data Dictionary"
    title = "Database Documentation"
    description = ""

    if os.path.exists(filepath):
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                lines = f.readlines()
                for line in lines:
                    if line.lower().startswith("project:"):
                        project = line.split(":", 1)[1].strip().upper()
                    elif line.lower().startswith("title:"):
                        title = line.split(":", 1)[1].strip()
                    elif line.lower().startswith("description:"):
                        description = line.split(":", 1)[1].strip()
                    elif not title and not description:
                        pass
        except Exception as e:
            print(f"Warning: Could not read body file: {e}")

    return project, title, description


def convert_file_to_pdf(input_path, output_dir):
    print(f"Converting {input_path} to PDF...")
    try:
        # Flags:
        # --headless: run without GUI
        # --convert-to pdf: target format
        # --outdir: output directory
        cmd = ['soffice', '--headless', '--convert-to', 'pdf', input_path, '--outdir', output_dir]
        result = subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.PIPE)

        stderr_output = result.stderr.decode()
        if stderr_output and "javaldx" not in stderr_output:
            print(f"LibreOffice output: {stderr_output}")

        filename = os.path.basename(input_path)
        pdf_filename = os.path.splitext(filename)[0] + ".pdf"
        pdf_path = os.path.join(output_dir, pdf_filename)

        if os.path.exists(pdf_path):
            print(f"Successfully converted to {pdf_path}")
            return pdf_path
        else:
            print("Error: PDF output file not found after conversion.")
            return None
    except subprocess.CalledProcessError as e:
        print(f"Error converting file to PDF: {e}")
        return None
    except FileNotFoundError:
        print("Error: 'soffice' command not found. Please install LibreOffice.")
        return None


def merge_pdfs(pdf_list, output_path):
    """
    Merges multiple PDFs into one using pdfunite.
    """
    print(f"Merging {len(pdf_list)} PDFs into {output_path}...")
    try:
        cmd = ['pdfunite'] + pdf_list + [output_path]
        subprocess.run(cmd, check=True)
        if os.path.exists(output_path):
            print(f"Successfully created merged PDF: {output_path}")
            return True
        return False
    except subprocess.CalledProcessError as e:
        print(f"Error merging PDFs: {e}")
        return False
    except FileNotFoundError:
        print("Error: 'pdfunite' command not found. Please install poppler-utils.")
        return False


def generate_er_diagram_pdf(connection_string, output_pdf_path, exclude_tables=None):
    """
    Generates ER diagram directly to PDF.
    """
    print(f"Generating ER diagram PDF...")
    try:
        # Ensure landscape mode for ER diagram (Left-to-Right)
        import eralchemy2.cst as cst
        cst.DOT_GRAPH_BEGINNING = r"""
      graph {
         graph [rankdir=LR, nodesep=1.0, ranksep=1.0, pad="0.5"];
         node [label="\N",
             shape=plaintext
         ];
         edge [color=gray50,
             minlen=2,
             style=dashed
         ];
"""
        render_er(connection_string, output_pdf_path, exclude_tables=exclude_tables)
        print(f"ER diagram PDF created: {output_pdf_path}")
        return True
    except Exception as e:
        print(f"Failed to generate ER diagram: {e}")
        return False


def create_intro_doc(model_path, title, description, output_doc_path, connection_string=None, exclude_tables=None):
    if os.path.exists(model_path):
        document = Document(model_path)
    else:
        print(f"Warning: Model file '{model_path}' not found. Using default blank document.")
        document = Document()

    # Clearing existing content to prevent duplication if model has placeholders
    if document.paragraphs:
        for p in list(document.paragraphs):
            p._element.getparent().remove(p._element)

    # Inserting title: Arial, Bold, 14, Centered
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run(title)
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(14)
    run_t.bold = True

    # Inserting description: Arial, 14, Justified
    p_desc = document.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_d = p_desc.add_run(description)
    run_d.font.name = 'Arial'
    run_d.font.size = Pt(14)

    document.save(output_doc_path)
    print(f"Intro document created: {output_doc_path}")


def generate_data_dictionary_xlsx(connection_string, output_xlsx_path, project_name=None, exclude_tables=None):
    """
    Generates Data Dictionary in XLSX format using SQLAlchemy for inspection.
    """
    print(f"Generating Data Dictionary XLSX...")
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from sqlalchemy import create_engine, inspect

        # Colors and Styles
        # New green color #B2D235
        header_fill = PatternFill(start_color="B2D235", end_color="B2D235", fill_type="solid")
        thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
        bold_font = Font(bold=True)
        center_align = Alignment(horizontal="center", vertical="center")

        wb = Workbook()
        ws = wb.active

        # Set Worksheet title
        ws.title = "Dicionário de Dados"

        # 1. Main Title - "Dicionário de dados", matching Intro Doc style (Arial, 14, Bold)
        ws.merge_cells('A2:I3')
        title_cell = ws['A2']
        title_cell.value = "Dicionário de dados"
        title_cell.font = Font(name='Arial', size=14, bold=True)
        title_cell.alignment = center_align
        # No fill, no border for the main title to match document style

        # 2. Database Inspection
        engine = create_engine(connection_string)
        insp = inspect(engine)
        all_tables = insp.get_table_names()

        # Filter exclusions
        tables_to_process = []
        if exclude_tables:
            tables_to_process = [t for t in all_tables if t not in exclude_tables]
        else:
            tables_to_process = all_tables

        tables_to_process.sort()

        current_row = 5  # Start after title

        for table_name in tables_to_process:
            # Table Name Header
            ws.merge_cells(f'A{current_row}:I{current_row}')
            cell_tn = ws[f'A{current_row}']
            cell_tn.value = table_name
            cell_tn.font = bold_font
            cell_tn.alignment = center_align
            cell_tn.fill = header_fill

            # Apply border to table name row
            for col in range(1, 10):
                ws.cell(row=current_row, column=col).border = thin_border

            current_row += 1

            # Column Headers
            headers = ["Campo", "Tipo", "Tamanho", "Precisão", "Obrigatório", "Único", "Chave primária", "Auto Incremento", "Chave estrangeira"]
            for idx, h in enumerate(headers, 1):
                cell = ws.cell(row=current_row, column=idx, value=h)
                cell.font = bold_font
                cell.alignment = center_align
                cell.fill = header_fill
                cell.border = thin_border

            current_row += 1

            # Get Columns, FKs, Unique Constraints
            columns = insp.get_columns(table_name)
            fks = insp.get_foreign_keys(table_name)
            unique_cons = insp.get_unique_constraints(table_name)
            indexes = insp.get_indexes(table_name)

            # Identify unique columns (single column constraints/indexes)
            unique_columns = set()
            for cons in unique_cons:
                if len(cons['column_names']) == 1:
                    unique_columns.add(cons['column_names'][0])
            for idx in indexes:
                if idx['unique'] and len(idx['column_names']) == 1:
                    unique_columns.add(idx['column_names'][0])

            # Map FKs: column_name -> referenced_table
            fk_map = {}
            for fk in fks:
                # constrained_columns is a list
                for local_col in fk['constrained_columns']:
                    fk_map[local_col] = fk['referred_table']

            for col in columns:
                c_name = col['name']
                c_type = col['type']

                # Determine attributes
                length = getattr(c_type, 'length', None)
                precision = getattr(c_type, 'precision', None)

                # Format checks
                val_tamanho = str(length) if length is not None else "-"
                val_precisao = str(precision) if precision is not None else "-"
                val_obrigatorio = "Sim" if not col.get('nullable', True) else "Não"
                val_unico = "Sim" if c_name in unique_columns else "Não"
                val_pk = "Sim" if col.get('primary_key') else "Não"
                val_auto = "Sim" if col.get('autoincrement') else "Não"
                val_fk = fk_map.get(c_name, "-")

                row_data = [
                    c_name,
                    str(c_type).split('(')[0],  # Simplified type name
                    val_tamanho,
                    val_precisao,
                    val_obrigatorio,
                    val_unico,
                    val_pk,
                    val_auto,
                    val_fk
                ]

                for idx, val in enumerate(row_data, 1):
                    cell = ws.cell(row=current_row, column=idx, value=val)
                    cell.alignment = Alignment(horizontal="center")
                    cell.border = thin_border

                current_row += 1

            # Empty row after each table
            current_row += 1

        # Adjust Column Widths
        column_widths = [20, 15, 10, 10, 12, 10, 15, 18, 20]
        for i, width in enumerate(column_widths, 1):
            ws.column_dimensions[chr(64 + i)].width = width

        # Print Settings: Fit to Width (1 page wide, n pages tall)
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 0  # 0 means auto/infinity
        ws.sheet_properties.pageSetUpPr.fitToPage = True
        # Optional: Landscape might be better for 9 columns
        ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
        # Margin narrowing (optional, using narrow margins if possible or defaults)
        from openpyxl.worksheet.page import PageMargins
        # ws.page_margins = PageMargins(left=0.25, right=0.25, top=0.75, bottom=0.75, header=0.3, footer=0.3)

        wb.save(output_xlsx_path)
        print(f"Data Dictionary XLSX created: {output_xlsx_path}")
        return True
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"Failed to generate Data Dictionary XLSX: {e}")
        return False


def main():
    print("=== Database Documentation Generator ===")

    script_dir = os.path.dirname(os.path.abspath(__file__))
    excluded_tables_path = os.path.join(script_dir, "excluded_tables.txt")
    body_path = os.path.join(script_dir, "details.txt")
    model_path = os.path.join(script_dir, "model.docx")

    host = input("Host (default: localhost): ").strip() or "localhost"
    default_port = '5432'
    port = input(f"Port (default: {default_port}): ").strip() or default_port
    user = input("Username: ").strip()
    password = input("Password: ").strip()
    db_name = input("Database Name: ").strip()

    if not all([user, password, db_name]):
        print("Error: Username, Password and Database Name are required.")
        return

    # Reading excluded tables
    exclude_tables = None
    if os.path.exists(excluded_tables_path):
        try:
            with open(excluded_tables_path, 'r', encoding='utf-8') as f:
                content = f.read().strip()
                if content:
                    exclude_tables = [t.strip() for t in content.split(',')]
            print(f"Loaded excluded tables: {len(exclude_tables)} tables.")
        except Exception as e:
            print(f"Warning: Could not read excluded_tables.txt: {e}")

    # Reading body content
    project_name, title, description = read_body_content(body_path)
    print(f"Project: {project_name}")
    print(f"Report Title: {title}")

    # Establishing output filenames
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    temp_er_pdf = os.path.join(script_dir, f"temp_er_{timestamp}.pdf")
    temp_intro_docx = os.path.join(script_dir, f"temp_intro_{timestamp}.docx")
    final_report_pdf = os.path.join(script_dir, f"Autodoc_{db_name}_{timestamp}.pdf")
    data_dict_xlsx = os.path.join(script_dir, f"Autodoc_DataDictionary_{timestamp}.xlsx")

    try:
        conn_str = create_connection_string(host, user, password, db_name, port)

        # 1. Generating ER Diagram PDF
        if not generate_er_diagram_pdf(conn_str, temp_er_pdf, exclude_tables=exclude_tables):
            print("Aborting: ER Diagram generation failed.")
            return

        # 1.2 Generate Data Dictionary XLSX
        data_dict_pdf = None
        if generate_data_dictionary_xlsx(conn_str, data_dict_xlsx, project_name=project_name, exclude_tables=exclude_tables):
            data_dict_pdf = convert_file_to_pdf(data_dict_xlsx, script_dir)

        # 2. Creating Intro DOCX
        create_intro_doc(model_path, title, description, temp_intro_docx)

        # 3. Converting Intro DOCX to PDF
        temp_intro_pdf = convert_file_to_pdf(temp_intro_docx, script_dir)

        if temp_intro_pdf and os.path.exists(temp_intro_pdf):
            # 4. Merging PDFs: Intro -> Data Dictionary -> ER Diagram
            pdf_list = [temp_intro_pdf]
            if data_dict_pdf and os.path.exists(data_dict_pdf):
                pdf_list.append(data_dict_pdf)
            pdf_list.append(temp_er_pdf)

            if merge_pdfs(pdf_list, final_report_pdf):
                print(f"\nSUCCESS! Final Report saved to: {final_report_pdf}")
            else:
                print("Failed to merge PDFs.")

            # Cleaning up
            try:
                os.remove(temp_intro_pdf)
                if data_dict_pdf:
                    os.remove(data_dict_pdf)
            except:
                pass
        else:
            print("Failed to convert intro document to PDF.")

    except Exception as e:
        print(f"An unexpected error occurred: {e}")
    finally:
        # Cleaning up temp files
        if os.path.exists(temp_er_pdf):
            os.remove(temp_er_pdf)
        if os.path.exists(temp_intro_docx):
            os.remove(temp_intro_docx)


if __name__ == "__main__":
    main()
